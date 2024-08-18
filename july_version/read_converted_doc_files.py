from docx import Document
from pdf2docx import Converter
from pdf2docx import parse
import os


# SRC = "docs1/smpcs/jylamvo-epar-product-information_en.docx"
# SRC = "docs1/smpcs/converted_to_docx/jylamvo-epar-product-information_en.docx"


# document = Document(SRC)


# def convert_pdfs_to_docx(pdfs):
#     for pdf in pdfs:
#         src = "docs1/smpcs/" + pdf
#         dst = "docs1/smpcs/converted_to_docx/" + pdf[:-3] + "docx"
#         # parse(src, dst)
#         cv = Converter(src)
#         cv.convert(dst, start=0, end=10)
#         cv.close()


def get_indications(document: Document):
    final_text = ""
    use = False
    start_header = "Terapeutiske indikationer"
    end_headers = ["Dosering og administration", "Dosering og indgivelsesmåde"]
    # paras = document.paragraphs
    # start = para.index("Therapeutic indications")
    # end = para.index("")
    for i, para in enumerate(document.paragraphs):
        if i == 100:
            break
        text = para.text
        if start_header in text:
            use = True
        elif any(end_header in text for end_header in end_headers):
            use = False
            break
        elif use:
            style_name = para.style.name
            if not text:
                final_text += "\n"
            elif style_name == "List Paragraph":
                final_text += "\t•\t" + text + "\n"
            else:
                final_text += text + "\n"
            # print(para.style.name)
            # print(text)
            # print("-----------")
    return final_text


def main():
    base_folder = "./docs1/doc_files/"
    for file in os.listdir(base_folder):
        if file.endswith(".docx"):
            src = os.path.join(base_folder, file)
            # print(src)
            document = Document(src)
            print("file: ", file)
            print("therapeutic indications text:\n\n")
            print(get_indications(document))
            print("----------------\n")
            # print(os.path.join("/mydir", file))


if __name__ == "__main__":
    main()

# final_text = get_indications(document)
# print(final_text)
# smpcs = [
#     "fortacin-epar-product-information_en.pdf",
#     "lacosamide-adroiq-epar-product-information_en.pdf",
#     "olanzapine-mylan-epar-product-information_en.pdf",
#     "ogluo-epar-product-information_en.pdf",
#     "jylamvo-epar-product-information_en.pdf",
#     "circadin-epar-product-information_en.pdf",
# ]
# convert_pdfs_to_docx(smpcs[4:5])

# for i, para in enumerate(document.paragraphs):
#     if "In rheumatological and dermatological diseases" in para.text:
#         runs = para.runs
#         for run in runs:
#             print(run.font.underline)
#         print("---------\n")
#     if "Rheumatological and dermatological diseases" in para.text:
#         runs = para.runs
#         for run in runs:
#             print("underline:\t", run.font.underline)
#             print("italic:\t\t", run.font.italic)
#         print("---------\n")
# if "Jylamvo is for use in the following indications:" in para.text:
#     runs = para.runs
#     for run in runs:
#         print(run.text)
#     print("---------\n")

# for i, para in enumerate(document.paragraphs):
#     # print(para)
#     if "Therapeutic indications" in para.text:
#         show = True
#     if show:
#         print(para.style.name)
#         print("Runs:")
#         for run in para.runs:
#             if run.italic:
#                 print("\titalic:", run.italic)
#             if run.style.name == "Emphasis":
#                 print("\tEmphasis:", run.style.name)
#         print(para.text)
#         print("---------\n")
#     if "Posology and method of administration" in para.text:
#         show = False
#         break
